#!/usr/bin/env python3
"""
preencher_relatorio_gui.py
Versão robusta do gerador de relatórios com fallback para CLI quando o Tkinter
não estiver disponível (resolve ModuleNotFoundError: No module named 'tkinter').
Funcionalidades:
- Consulta ReceitaWS por CNPJ
- Preenche placeholders em template .docx
- Gera [OBJETIVO_EMPRESA] opcional via provedor de IA (pluggable)
- Insere hyperlink para [LINK_DRIVE] e [LINK_PARA_DOWNLOAD]
- Insere imagem para [IDENTIDADE_VISUAL_E_PALETA_DE_CORES] e páginas específicas
- Modo GUI (Tkinter) quando disponível; caso contrário, modo CLI automático
- Argumentos de linha de comando para rodar em modo não-GUI
- Testes unitários simples acessíveis via --run-tests
"""
from __future__ import annotations
import re
import os
import time
import sys
import json
import argparse
import requests
import shutil
import tempfile
from pathlib import Path
from typing import Dict, Optional

# Selenium (opcional)
try:
    from selenium import webdriver
    from selenium.webdriver.common.By import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options
    from selenium.common.exceptions import TimeoutException, NoSuchElementException, WebDriverException
    SELENIUM_AVAILABLE = True
except Exception:
    SELENIUM_AVAILABLE = False
from docx.shared import Inches
# tenta importar tkinter dinamicamente (alguns ambientes não têm suporte)
try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
    TKINTER_AVAILABLE = True
except Exception:
    TKINTER_AVAILABLE = False
    tk = None
    filedialog = None
    messagebox = None
# docx (necessário instalar python-docx)
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Pillow (opcional para colar imagens)
try:
    from PIL import ImageGrab, Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    ImageGrab = None
    Image = None


# ----------------- Configuração -----------------
RECEITAWS_URL = "https://www.receitaws.com.br/v1/cnpj/{}"
REQUEST_TIMEOUT = 10
PLACEHOLDER_PATTERN = re.compile(r'\[([A-Z0-9_]+)\]')
# Lista de campos de imagem
IMAGE_FIELDS = [
    "IDENTIDADE_VISUAL_E_PALETA_DE_CORES",
    "PAGINA_HOME_IMG",
    "PAGINA_PRODUTOS_IMG",
    "PAGINA_QUEM_SOMOS_IMG",
    "PAGINA_CONTATO_IMG",
    "DETALHES_PEDIDO",
    "DETALHES_PRODUTO",
    "TODOS_PRODUTOS"
]
# ----------------- Utilitários -----------------
def normalize_cnpj(cnpj_raw: str) -> str:
    digits = re.sub(r'\D', '', cnpj_raw or '')
    if len(digits) != 14:
        raise ValueError("CNPJ deve conter 14 dígitos (após remover pontuação).")
    return digits

def normalize_image_path(image_path: str) -> str:
    """
    Normaliza o caminho da imagem para evitar problemas com espaços.
    Se o caminho contém espaços, cria uma cópia temporária sem espaços.
    """
    if not image_path or not Path(image_path).exists():
        return image_path
    # Se não tem espaços, retorna o caminho original
    if ' ' not in image_path:
        return image_path
    # Cria um arquivo temporário sem espaços
    original_path = Path(image_path)
    file_extension = original_path.suffix
    temp_dir = Path(tempfile.gettempdir()) / "docx_images"
    temp_dir.mkdir(exist_ok=True)
    # Gera nome sem espaços baseado no nome original
    safe_name = re.sub(r'[^\w\-_.]', '_', original_path.stem)
    temp_path = temp_dir / f"{safe_name}{file_extension}"
    # Copia o arquivo para o local temporário
    try:
        shutil.copy2(image_path, temp_path)
        return str(temp_path)
    except Exception as e:
        print(f"Warning: Não foi possível criar cópia temporária da imagem: {e}", file=sys.stderr)
        return image_path

def consulta_empresa(cnpj: str) -> dict:
    url = RECEITAWS_URL.format(cnpj)
    tries = 0
    while tries < 3:
        tries += 1
        try:
            resp = requests.get(url, timeout=REQUEST_TIMEOUT)
            if resp.status_code != 200:
                if resp.status_code in (429, 500, 502, 503, 504):
                    time.sleep(1 + tries)
                    continue
                resp.raise_for_status()
            data = resp.json()
            if isinstance(data, dict) and data.get("status") == "ERROR":
                raise RuntimeError(f"ReceitaWS retornou erro: {data.get('message')}")
            return data
        except requests.RequestException as e:
            if tries >= 3:
                raise RuntimeError(f"Falha ao consultar ReceitaWS: {e}")
            time.sleep(1 + tries)
    raise RuntimeError("Não foi possível consultar ReceitaWS após tentativas.")

def build_mapping(data: dict) -> dict:
    def safe_get(key, default=""):
        v = data.get(key)
        if v is None:
            return default
        if isinstance(v, str):
            return v.strip()
        return str(v)

    atividade_principal = ""
    if data.get("atividade_principal"):
        try:
            atividade_principal = data["atividade_principal"][0].get("text", "")
        except Exception:
            atividade_principal = str(data.get("atividade_principal"))

    resumo = " | ".join(filter(None, [
        atividade_principal,
        safe_get("porte"),
        safe_get("situacao")
    ]))

    endereco = " - ".join(filter(None, [
        safe_get("logradouro"),
        safe_get("numero"),
        safe_get("bairro"),
        safe_get("municipio"),
        safe_get("uf"),
        safe_get("cep"),
    ]))

    mapping = {
        "NOME_EMPRESA_CLIENTE": safe_get("nome"),
        "FANTASIA": safe_get("fantasia"),
        "RESUMO_EMPRESA_CLIENTE": resumo,
        "CNPJ": safe_get("cnpj"),
        "ENDERECO": endereco,
        "ATIVIDADE_PRINCIPAL": atividade_principal,
        "TELEFONE": safe_get("telefone"),
        "EMAIL": safe_get("email"),
        "ABERTURA": safe_get("abertura"),
        "SITUACAO": safe_get("situacao"),
        "OBJETIVO_EMPRESA": "",
        "LINK_DRIVE": "",
        "LINK_DRIVE_TEXT": "",
        "LINK_PARA_DOWNLOAD": "",
        "LINK_PARA_DOWNLOAD_TEXT": "",
        "DATA_BACKUP": "",
        "DATA_KICKOFF": "",
        "DATA_ENTREGA": "",
        "DOMINIO": "",
        "DEMANDA": "",
        "DOMINIOWP": "",
        "ESPECIALISTARESPONSAVEL": "ITALO FELIPE IGNACIO",
        "IDENTIDADE_VISUAL_E_PALETA_DE_CORES": "",
        "PAGINA_HOME_IMG": "",
        "PAGINA_PRODUTOS_IMG": "",
        "PAGINA_QUEM_SOMOS_IMG": "",
        "PAGINA_CONTATO_IMG": "",
        "HOSPEDAGEM": "",
        "EMAIL_CRIADO": "",
        "DETALHES_PEDIDO": "",
        "DETALHES_PRODUTO": "",
        "TODOS_PRODUTOS": "",
    }
    return mapping

# ----------------- AI Providers -----------------
class AIProviderBase:
    def generate_objective(self, source_text: str, context: dict) -> str:
        raise NotImplementedError

class MockProvider(AIProviderBase):
    def generate_objective(self, source_text: str, context: dict) -> str:
        nome = context.get("NOME_EMPRESA_CLIENTE", "").strip()
        atividade = context.get("ATIVIDADE_PRINCIPAL", "").strip()
        if atividade:
            return (
                f"O objetivo da {nome} é atuar em {atividade.lower()}, oferecendo soluções "
                "e serviços relacionados a essa atividade, com foco em qualidade e atendimento ao cliente."
            )
        if source_text:
            first_sent = source_text.split(".")[0].strip()
            if first_sent:
                return f"O objetivo da {nome} é {first_sent}."
        return f"O objetivo da {nome} é oferecer produtos/serviços no seu segmento de atuação."

class HuggingFaceProvider(AIProviderBase):
    def __init__(self, api_token: Optional[str] = None, model: str = "google/flan-t5-large"):
        self.api_token = api_token or os.environ.get("HUGGINGFACE_API_TOKEN")
        self.model = model
        if not self.api_token:
            raise RuntimeError("Hugging Face token não configurado (HUGGINGFACE_API_TOKEN).")

    def generate_objective(self, source_text: str, context: dict) -> str:
        prompt = (
            "Você é um assistente que escreve um 'Objetivo da Empresa' curto (1-2 parágrafos) "
            "baseado nas informações abaixo. Seja direto e formal.\n\n"
            f"INFORMAÇÕES:\n{source_text}\n\n"
            "RETORNE APENAS o texto final, sem rótulos."
        )
        url = f"https://api-inference.huggingface.co/models/{self.model}"
        headers = {"Authorization": f"Bearer {self.api_token}"}
        payload = {"inputs": prompt, "options": {"wait_for_model": True}}
        resp = requests.post(url, json=payload, headers=headers, timeout=30)
        if resp.status_code != 200:
            raise RuntimeError(f"HF API erro {resp.status_code}: {resp.text}")
        result = resp.json()
        text = ""
        if isinstance(result, list) and result:
            first = result[0]
            if isinstance(first, dict):
                text = first.get("generated_text") or first.get("text") or str(first)
            else:
                text = str(first)
        elif isinstance(result, dict):
            text = result.get("generated_text") or result.get("text") or json.dumps(result)
        else:
            text = str(result)
        return (text or "").strip()

class OpenAIProvider(AIProviderBase):
    def __init__(self, api_key: Optional[str] = None, model: str = "gpt-4o-mini"):
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY")
        self.model = model
        if not self.api_key:
            raise RuntimeError("OPENAI_API_KEY não configurada.")
        try:
            from openai import OpenAI
            self.client = OpenAI(api_key=self.api_key)
        except Exception as e:
            raise RuntimeError("Biblioteca openai não instalada. pip install openai") from e

    def generate_objective(self, source_text: str, context: dict) -> str:
        prompt = (
            "Escreva um texto curto (1-2 parágrafos) intitulado 'Objetivo da Empresa' baseado nas "
            "informações abaixo. Use linguagem formal e direta. Retorne apenas o texto.\n\n"
            f"INFORMAÇÕES:\n{source_text}\n"
        )
        resp = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=256,
            temperature=0.2,
        )
        content = resp.choices[0].message.content.strip()
        return content

def get_ai_provider(name: Optional[str]) -> AIProviderBase:
    name = (name or "mock").lower()
    if name == "mock":
        return MockProvider()
    if name in ("hf", "huggingface"):
        return HuggingFaceProvider()
    if name in ("openai", "gpt"):
        return OpenAIProvider()
    raise ValueError(f"Provider IA desconhecido: {name}")

# ----------------- DOCX helpers -----------------
def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_run.append(new_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def replace_in_paragraph(paragraph, mapping: Dict[str, str]):
    full_text = "".join([r.text for r in paragraph.runs])
    non_link_keys = {k: v for k, v in mapping.items() if k not in ("LINK_DRIVE", "LINK_DRIVE_TEXT", "LINK_PARA_DOWNLOAD", "LINK_PARA_DOWNLOAD_TEXT") + tuple(IMAGE_FIELDS)}
    # Handle [LINK_DRIVE]
    if "[LINK_DRIVE]" in full_text and mapping.get("LINK_DRIVE"):
        parts = full_text.split("[LINK_DRIVE]")
        # Clear all runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)
        for idx, part in enumerate(parts):
            # Replace other placeholders in this part
            for key, val in non_link_keys.items():
                part = part.replace(f'[{key}]', val or "")
            if part:
                paragraph.add_run(part)
            if idx < len(parts) - 1:
                display = mapping.get("LINK_DRIVE_TEXT") or "Link Drive"
                add_hyperlink(paragraph, mapping["LINK_DRIVE"], display)
        return
    # Handle [LINK_PARA_DOWNLOAD]
    if "[LINK_PARA_DOWNLOAD]" in full_text and mapping.get("LINK_PARA_DOWNLOAD"):
        parts = full_text.split("[LINK_PARA_DOWNLOAD]")
        # Clear all runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)
        for idx, part in enumerate(parts):
            # Replace other placeholders in this part
            for key, val in non_link_keys.items():
                part = part.replace(f'[{key}]', val or "")
            if part:
                paragraph.add_run(part)
            if idx < len(parts) - 1:
                display = "Link para download"
                add_hyperlink(paragraph, mapping["LINK_PARA_DOWNLOAD"], display)
        return
    # Handle image fields
    for image_field in IMAGE_FIELDS:
        if f"[{image_field}]" in full_text:
            image_path = mapping.get(image_field)
            # Clear all runs
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            if image_path:
                # Normaliza o caminho para evitar problemas com espaços
                normalized_path = normalize_image_path(image_path)
                if Path(normalized_path).exists():
                    try:
                        run = paragraph.add_run()
                        run.add_picture(normalized_path, width=Inches(5.0))
                    except Exception as e:
                        error_msg = f"Erro ao inserir imagem {image_field}: {e}"
                        print(error_msg, file=sys.stderr)
                        paragraph.add_run(error_msg)
                else:
                    paragraph.add_run(f"Imagem não encontrada para {image_field}: {image_path}")
            else:
                paragraph.add_run(f"Nenhuma imagem fornecida para {image_field}.")
            return
    # Normal case: per-run replacement to preserve formatting
    for run in paragraph.runs:
        text = run.text
        for key, val in non_link_keys.items():
            text = text.replace(f'[{key}]', val or "")
        run.text = text
    # Check for remaining placeholders (spanning runs)
    new_full_text = "".join([r.text for r in paragraph.runs])
    remaining = PLACEHOLDER_PATTERN.findall(new_full_text)
    if remaining:
        # Fallback: rebuild with single run (loses formatting for spanning parts)
        new_text = new_full_text
        for key, val in non_link_keys.items():
            new_text = new_text.replace(f'[{key}]', val or "")
        # Clear runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)
        # Add new run
        paragraph.add_run(new_text)
        print(f"Warning: Rebuilt paragraph due to spanning placeholders: {remaining}", file=sys.stderr)

def replace_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            replace_in_block(cell, mapping)

def replace_in_block(block, mapping: Dict[str, str]):
    for para in block.paragraphs:
        replace_in_paragraph(para, mapping)
    for table in getattr(block, "tables", []):
        replace_in_table(table, mapping)

def process_document(template_path: str, output_path: str, mapping: Dict[str, str]):
    doc = Document(template_path)
    replace_in_block(doc, mapping)
    for section in doc.sections:
        if section.header:
            replace_in_block(section.header, mapping)
        if section.footer:
            replace_in_block(section.footer, mapping)
    doc.save(output_path)

def fix_docx_templates():
    """Fix para PyInstaller não incluir templates docx"""
    import sys
    import os
    if hasattr(sys, '_MEIPASS'):
        # Estamos executando via PyInstaller
        import docx
        template_dir = os.path.join(sys._MEIPASS, 'docx', 'templates')
        if os.path.exists(template_dir):
            docx.shared.TEMPLATE_DIR = template_dir

# ----------------- CLI flow -----------------
def run_cli(template: Optional[str] = None, cnpj: Optional[str] = None, drive: Optional[str] = None,
            drive_text: Optional[str] = None, use_ai: Optional[bool] = None, ai_provider: Optional[str] = None,
            out: Optional[str] = None, extra_mapping: Optional[dict] = None) -> None:
    try:
        if not template:
            template = input("Caminho do template .docx: ").strip()
        if not Path(template).exists():
            print("Template não encontrado:", template)
            return
        if not cnpj:
            cnpj = input("CNPJ da empresa: ").strip()
        try:
            cnpj_norm = normalize_cnpj(cnpj)
        except Exception as e:
            print("CNPJ inválido:", e)
            return
        print("Consultando ReceitaWS...")
        data = consulta_empresa(cnpj_norm)
        mapping = build_mapping(data)
        if drive is None:
            drive = input("Link do Drive (opcional, ENTER para pular): ").strip()
        if drive:
            if not drive.startswith(("http://", "https://")):
                drive = "https://" + drive
            mapping["LINK_DRIVE"] = drive
            mapping["LINK_DRIVE_TEXT"] = drive_text or input("Texto do link (ENTER para 'Link Drive'): ").strip() or "Link Drive"
            mapping["LINK_PARA_DOWNLOAD"] = drive
            mapping["LINK_PARA_DOWNLOAD_TEXT"] = "Link para download"
        # Handle extra fields
        extra_fields = {
            "DATA_BACKUP": "Data Backup (opcional): ",
            "DATA_KICKOFF": "Data Kickoff (opcional): ",
            "DATA_ENTREGA": "Data Entrega (opcional): ",
            "DOMINIO": "Domínio (opcional): ",
            "DEMANDA": "Demanda (opcional): ",
            "HOSPEDAGEM": "Hospedagem (HOSTINGER, LOCALWEB, UOL ou custom, separados por vírgula): ",
            "EMAIL_CRIADO": "Email Criado (opcional): ",
            "IDENTIDADE_VISUAL_E_PALETA_DE_CORES": "Caminho da imagem para Identidade Visual e Paleta de Cores (opcional): ",
            "PAGINA_HOME_IMG": "Caminho da imagem para Página Home (opcional): ",
            "PAGINA_PRODUTOS_IMG": "Caminho da imagem para Página Produtos (opcional): ",
            "PAGINA_QUEM_SOMOS_IMG": "Caminho da imagem para Página Quem Somos (opcional): ",
            "PAGINA_CONTATO_IMG": "Caminho da imagem para Página Contato (opcional): ",
            "DETALHES_PEDIDO": "Caminho da imagem para Detalhes Pedido (opcional): ",
            "DETALHES_PRODUTO": "Caminho da imagem para Detalhes Produto (opcional): ",
            "TODOS_PRODUTOS": "Caminho da imagem para Todos Produtos (opcional): ",
        }
        for field, prompt in extra_fields.items():
            value = extra_mapping.get(field, "") if extra_mapping else ""
            if not value:
                value = input(prompt).strip()
            mapping[field] = value
        # Auto-set DOMINIOWP based on DOMINIO
        if mapping.get("DOMINIO"):
            mapping["DOMINIOWP"] = mapping["DOMINIO"] + "/wp-admin/"
        else:
            mapping["DOMINIOWP"] = ""
        # ESPECIALISTARESPONSAVEL is always fixed
        mapping["ESPECIALISTARESPONSAVEL"] = "ITALO FELIPE IGNACIO"
        if use_ai is None:
            use_ai = input("Deseja usar IA para preencher [OBJETIVO_EMPRESA]? (s/N): ").strip().lower() == 's'
        if use_ai:
            provider = (ai_provider or os.environ.get("AI_PROVIDER") or input("Provedor IA (mock/hf/openai) [mock]: ").strip() or "mock")
            try:
                ai = get_ai_provider(provider)
            except Exception as e:
                print("Erro ao inicializar provedor IA:", e)
                print("Usando MockProvider como fallback.")
                ai = MockProvider()
            source_parts = []
            if mapping.get("ATIVIDADE_PRINCIPAL"):
                source_parts.append("Atividade principal: " + mapping["ATIVIDADE_PRINCIPAL"])
            if mapping.get("RESUMO_EMPRESA_CLIENTE"):
                source_parts.append("Resumo: " + mapping["RESUMO_EMPRESA_CLIENTE"])
            source_text = "\n".join(source_parts).strip() or str(data)[:2000]
            try:
                mapping["OBJETIVO_EMPRESA"] = ai.generate_objective(source_text, mapping)
            except Exception as e:
                print("Erro ao gerar objetivo com IA:", e)
                print("Usando heurística local.")
                mapping["OBJETIVO_EMPRESA"] = MockProvider().generate_objective(source_text, mapping)
        else:
            mapping["OBJETIVO_EMPRESA"] = ""
        out_path = out or input("Arquivo de saída (.docx) [relatorio_saida.docx]: ").strip() or f'relatorio_{cnpj_norm}.docx'
        print("Gerando documento...")
        process_document(template, out_path, mapping)
        print("Documento gerado:", out_path)
    except Exception as e:
        print("Erro durante execução:", e)

# ----------------- GUI flow -----------------
if TKINTER_AVAILABLE:
    class App:
        def __init__(self, root):
            self.root = root
            root.title("Gerador de Relatório - CNPJ -> Word")
            # Criar frame principal com scrollbar
            main_frame = tk.Frame(root)
            main_frame.pack(fill=tk.BOTH, expand=True)
            canvas = tk.Canvas(main_frame)
            scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas)
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            frm = tk.Frame(scrollable_frame, padx=10, pady=10)
            frm.pack(fill=tk.BOTH, expand=True)
            row = 0
            # Template
            tk.Label(frm, text="Template (.docx):").grid(row=row, column=0, sticky='w')
            self.entry_template = tk.Entry(frm, width=60)
            self.entry_template.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Abrir", command=self.browse_template).grid(row=row, column=2)
            row += 1
            # CNPJ
            tk.Label(frm, text="CNPJ da empresa:").grid(row=row, column=0, sticky='w', pady=(10,0))
            self.entry_cnpj = tk.Entry(frm, width=40)
            self.entry_cnpj.grid(row=row, column=1, sticky='w', pady=(10,0))
            row += 1
            # Link Drive
            tk.Label(frm, text="Link do Drive (opcional):").grid(row=row, column=0, sticky='w')
            self.entry_drive = tk.Entry(frm, width=60)
            self.entry_drive.grid(row=row, column=1, sticky='w')
            tk.Label(frm, text="Texto do Link:").grid(row=row, column=2, sticky='w')
            self.entry_drive_text = tk.Entry(frm, width=20)
            self.entry_drive_text.grid(row=row, column=3, sticky='w')
            row += 1
            # Extra fields
            tk.Label(frm, text="Data Backup:").grid(row=row, column=0, sticky='w')
            self.entry_data_backup = tk.Entry(frm, width=20)
            self.entry_data_backup.grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Data Kickoff:").grid(row=row, column=0, sticky='w')
            self.entry_data_kickoff = tk.Entry(frm, width=20)
            self.entry_data_kickoff.grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Data Entrega:").grid(row=row, column=0, sticky='w')
            self.entry_data_entrega = tk.Entry(frm, width=20)
            self.entry_data_entrega.grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Domínio:").grid(row=row, column=0, sticky='w')
            self.entry_dominio = tk.Entry(frm, width=40)
            self.entry_dominio.grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Demanda:").grid(row=row, column=0, sticky='w')
            self.entry_demanda = tk.Entry(frm, width=40)
            self.entry_demanda.grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Email Criado:").grid(row=row, column=0, sticky='w')
            self.entry_email_criado = tk.Entry(frm, width=40)
            self.entry_email_criado.grid(row=row, column=1, sticky='w')
            row += 1
            # Hospedagem
            tk.Label(frm, text="Hospedagem:").grid(row=row, column=0, sticky='w', pady=(10,0))
            self.hostinger_var = tk.IntVar()
            tk.Checkbutton(frm, text="HOSTINGER", variable=self.hostinger_var).grid(row=row, column=1, sticky='w')
            row += 1
            self.localweb_var = tk.IntVar()
            tk.Checkbutton(frm, text="LOCALWEB", variable=self.localweb_var).grid(row=row, column=1, sticky='w')
            row += 1
            self.uol_var = tk.IntVar()
            tk.Checkbutton(frm, text="UOL", variable=self.uol_var).grid(row=row, column=1, sticky='w')
            row += 1
            tk.Label(frm, text="Hospedagem Custom:").grid(row=row, column=0, sticky='w')
            self.entry_hospedagem_custom = tk.Entry(frm, width=40)
            self.entry_hospedagem_custom.grid(row=row, column=1, sticky='w')
            row += 1
            # Separador para imagens
            tk.Label(frm, text="--- IMAGENS ---", font=("Arial", 10, "bold")).grid(row=row, column=0, columnspan=4, pady=10)
            row += 1
            # Identidade Visual e Paleta de cores
            tk.Label(frm, text="Identidade Visual e Paleta de cores:").grid(row=row, column=0, sticky='w')
            self.entry_identidade = tk.Entry(frm, width=60)
            self.entry_identidade.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_identidade)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_identidade)).grid(row=row, column=3)
            row += 1
            # Página Home
            tk.Label(frm, text="Página Home:").grid(row=row, column=0, sticky='w')
            self.entry_pagina_home = tk.Entry(frm, width=60)
            self.entry_pagina_home.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_pagina_home)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_pagina_home)).grid(row=row, column=3)
            row += 1
            # Página Produtos
            tk.Label(frm, text="Página Produtos:").grid(row=row, column=0, sticky='w')
            self.entry_pagina_produtos = tk.Entry(frm, width=60)
            self.entry_pagina_produtos.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_pagina_produtos)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_pagina_produtos)).grid(row=row, column=3)
            row += 1
            # Página Quem Somos
            tk.Label(frm, text="Página Quem Somos:").grid(row=row, column=0, sticky='w')
            self.entry_pagina_quem_somos = tk.Entry(frm, width=60)
            self.entry_pagina_quem_somos.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_pagina_quem_somos)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_pagina_quem_somos)).grid(row=row, column=3)
            row += 1
            # Página Contato
            tk.Label(frm, text="Página Contato:").grid(row=row, column=0, sticky='w')
            self.entry_pagina_contato = tk.Entry(frm, width=60)
            self.entry_pagina_contato.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_pagina_contato)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_pagina_contato)).grid(row=row, column=3)
            row += 1
            # Detalhes Pedido
            tk.Label(frm, text="Detalhes Pedido:").grid(row=row, column=0, sticky='w')
            self.entry_detalhes_pedido = tk.Entry(frm, width=60)
            self.entry_detalhes_pedido.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_detalhes_pedido)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_detalhes_pedido)).grid(row=row, column=3)
            row += 1
            # Detalhes Produto
            tk.Label(frm, text="Detalhes Produto:").grid(row=row, column=0, sticky='w')
            self.entry_detalhes_produto = tk.Entry(frm, width=60)
            self.entry_detalhes_produto.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_detalhes_produto)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_detalhes_produto)).grid(row=row, column=3)
            row += 1
            # Todos Produtos
            tk.Label(frm, text="Todos Produtos:").grid(row=row, column=0, sticky='w')
            self.entry_todos_produtos = tk.Entry(frm, width=60)
            self.entry_todos_produtos.grid(row=row, column=1, sticky='w')
            tk.Button(frm, text="Selecionar", command=lambda: self.browse_image(self.entry_todos_produtos)).grid(row=row, column=2)
            tk.Button(frm, text="Colar", command=lambda: self.paste_image_from_clipboard(self.entry_todos_produtos)).grid(row=row, column=3)
            row += 1
            # Separador
            tk.Label(frm, text="--- CONFIGURAÇÕES ---", font=("Arial", 10, "bold")).grid(row=row, column=0, columnspan=3, pady=10)
            row += 1
            # Especialista
            tk.Label(frm, text="Especialista Responsável:").grid(row=row, column=0, sticky='w')
            self.entry_especialista = tk.Entry(frm, width=40)
            self.entry_especialista.grid(row=row, column=1, sticky='w')
            self.entry_especialista.insert(0, "ITALO FELIPE IGNACIO")
            self.entry_especialista.config(state="readonly") # Make it read-only
            row += 1
            # IA
            self.use_ai_var = tk.IntVar(value=1)
            tk.Checkbutton(frm, text="Usar IA para preencher [OBJETIVO_EMPRESA]", variable=self.use_ai_var).grid(row=row, column=0, sticky='w', columnspan=2, pady=(10,0))
            row += 1
            tk.Label(frm, text="Provedor IA:").grid(row=row, column=0, sticky='w', pady=(10,0))
            self.ai_provider = tk.StringVar(value=os.environ.get('AI_PROVIDER', 'mock'))
            tk.OptionMenu(frm, self.ai_provider, 'mock', 'hf', 'openai').grid(row=row, column=1, sticky='w')
            row += 1
            # Arquivo saída
            tk.Label(frm, text="Arquivo saída (.docx):").grid(row=row, column=0, sticky='w', pady=(10,0))
            self.entry_out = tk.Entry(frm, width=60)
            self.entry_out.grid(row=row, column=1, sticky='w')
            self.entry_out.insert(0, 'relatorio_saida.docx')
            row += 1
            # Botões de ação
            actions = tk.Frame(frm, pady=10)
            actions.grid(row=row, column=0, columnspan=3, sticky='w')
            tk.Button(actions, text="Gerar Relatório", command=self.run, bg="lightgreen", font=("Arial", 12, "bold")).grid(row=0, column=0, padx=(0,10))
            tk.Button(actions, text="Preencher automaticamente", command=self.auto_fill).grid(row=0, column=1)

        def browse_template(self):
            p = filedialog.askopenfilename(filetypes=[('Word files', '*.docx')])
            if p:
                self.entry_template.delete(0, tk.END)
                self.entry_template.insert(0, p)

        def browse_image(self, entry_widget):
            p = filedialog.askopenfilename(filetypes=[('Image files', '*.jpg *.jpeg *.png *.gif *.bmp *.tiff')])
            if p:
                entry_widget.delete(0, tk.END)
                entry_widget.insert(0, p)

        def paste_image_from_clipboard(self, entry_widget):
            """Tenta colar uma imagem da área de transferência no widget de entrada."""
            if not PILLOW_AVAILABLE:
                messagebox.showwarning("Funcionalidade Indisponível",
                                     "A biblioteca 'Pillow' é necessária para colar imagens. "
                                     "Instale com: pip install Pillow")
                return

            image = None
            try:
                # ImageGrab.grabclipboard() pode ser uma imagem ou uma lista de arquivos
                image = ImageGrab.grabclipboard()
            except Exception:
                # Se o clipboard não contiver uma imagem, o ImageGrab pode falhar.
                # Tentamos obter o conteúdo como texto, que pode ser um caminho de arquivo.
                pass

            # Se nenhuma imagem foi obtida, verifique se há um caminho de arquivo no clipboard
            if image is None:
                try:
                    clipboard_content = self.root.clipboard_get()
                    path = Path(clipboard_content.strip())
                    if path.exists() and path.is_file():
                        entry_widget.delete(0, tk.END)
                        entry_widget.insert(0, str(path.resolve()))
                        return
                except (tk.TclError, Exception):
                     # Não é um texto ou não é um caminho válido, ignorar.
                     pass
                
                messagebox.showinfo("Informação", "Nenhuma imagem ou caminho de arquivo válido encontrado na área de transferência.")
                return

            # Se for uma lista de arquivos (comum no Windows)
            if isinstance(image, list):
                filepath = image[0]
                if Path(filepath).exists():
                    entry_widget.delete(0, tk.END)
                    entry_widget.insert(0, str(Path(filepath).resolve()))
                    return
            
            # Se for um objeto de imagem
            if isinstance(image, Image.Image):
                try:
                    temp_dir = Path(tempfile.gettempdir()) / "docx_images_pasted"
                    temp_dir.mkdir(exist_ok=True)
                    
                    timestamp = int(time.time())
                    temp_path = temp_dir / f"pasted_image_{timestamp}.png"
                    
                    image.save(temp_path, "PNG")
                    
                    entry_widget.delete(0, tk.END)
                    entry_widget.insert(0, str(temp_path))
                except Exception as e:
                    messagebox.showerror("Erro", f"Falha ao salvar a imagem colada: {e}")
                return

            messagebox.showinfo("Informação", "O conteúdo da área de transferência não é uma imagem reconhecida.")

        def auto_fill(self):
            try:
                from selenium import webdriver
                from selenium.webdriver.common.by import By
                from selenium.webdriver.support.ui import WebDriverWait
                from selenium.webdriver.support import expected_conditions as EC
                from selenium.webdriver.chrome.options import Options
                from selenium.common.exceptions import TimeoutException, NoSuchElementException, WebDriverException
            except Exception as e:
                messagebox.showerror('Erro', f'Selenium não está disponível neste executável: {e}')
                return

            dominio = self.entry_dominio.get().strip()
            if not dominio:
                messagebox.showerror('Erro', 'Informe go Domínio (ex: exemplo.com)')
                return
            if not dominio.startswith(('http://', 'https://')):
                dominio = 'https://' + dominio
            url = dominio.rstrip('/') + '/wp-admin/'
            out_img = self.entry_identidade.get().strip()
            if not out_img:
                try:
                    base_dir = os.path.expanduser('~/Pictures/relatorios_auto')
                    os.makedirs(base_dir, exist_ok=True)
                except Exception:
                    base_dir = os.getcwd()
                domain_slug = re.sub(r'[^a-zA-Z0-9]+', '_', dominio)
                filename = f"s/style_guide_{domain_slug}_{int(time.time())}.png"
                out_img = os.path.join(base_dir, filename)
                self.entry_identidade.delete(0, tk.END)
                self.entry_identidade.insert(0, out_img)

            chrome_opts = Options()
            chrome_opts.add_argument('--no-sandbox')
            chrome_opts.add_argument('--disable-dev-shm-usage')
            chrome_opts.add_argument('--start-maximized')
            driver = None
            try:
                driver = webdriver.Chrome(options=chrome_opts)
                driver.set_page_load_timeout(30)
                driver.get(url)
                try:
                    driver.maximize_window()
                except Exception:
                    pass
                WebDriverWait(driver, 10).until(lambda d: d.execute_script('return document.readyState') == 'complete')

                # Login se necessário
                try:
                    user = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.ID, 'user_login')))
                    pwd = driver.find_element(By.ID, 'user_pass')
                    btn = driver.find_element(By.ID, 'wp-submit')
                    user.send_keys('admin')
                    pwd.send_keys('Senai@127')
                    btn.click()
                    WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'wpadminbar')))
                except TimeoutException:
                    pass

                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'adminmenumain')))

                # ----------------- FLUXO 1: Prints do site público -----------------
                print('[auto_fill] Iniciando fluxo de prints do site público...')
                try:
                    # Link do site no admin bar (href == domínio)
                    site_link = None
                    hrefs_to_try = [
                        dominio.rstrip('/'),
                        dominio.rstrip('/') + '/',
                        (('http://' + dominio.split('://')[-1]).rstrip('/')),
                        (('https://' + dominio.split('://')[-1]).rstrip('/')),
                    ]
                    for href_try in hrefs_to_try:
                        try:
                            site_link = WebDriverWait(driver, 10).until(
                                EC.element_to_be_clickable((By.CSS_SELECTOR, f'a.ab-item[href^="{href_try}"]'))
                            )
                            if site_link:
                                break
                        except TimeoutException:
                            continue
                    if not site_link:
                        # Fallback: procurar primeiro ab-item que aponte para root do domínio
                        links = driver.find_elements(By.CSS_SELECTOR, 'a.ab-item')
                        for lk in links:
                            try:
                                href = (lk.get_attribute('href') or '').rstrip('/')
                                if any(href.startswith(h) for h in hrefs_to_try):
                                    site_link = lk
                                    break
                            except Exception:
                                continue
                    if not site_link:
                        raise RuntimeError('Link do site no admin bar não encontrado.')
                    driver.execute_script('arguments[0].click();', site_link)
                except Exception as e:
                    print(f'[auto_fill] Falha ao navegar para o site público: {e}')
                    raise

                # Aguarda home carregar e tira screenshot -> Página Home
                try:
                    WebDriverWait(driver, 15).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, 'body, html'))
                    )
                except TimeoutException:
                    pass
                home_dir = os.path.dirname(out_img)
                domain_slug = re.sub(r'[^a-zA-Z0-9]+', '_', dominio)
                home_path = os.path.join(home_dir, f'home_{domain_slug}_{int(time.time())}.png')
                driver.save_screenshot(home_path)
                try:
                    self.entry_pagina_home.delete(0, tk.END)
                    self.entry_pagina_home.insert(0, home_path)
                except Exception:
                    pass
                print(f'[auto_fill] Home screenshot: {home_path}')

                # Helper para navegar por menus e printar
                def go_and_shoot(link_text_candidates, out_filename_prefix, gui_entry_widget):
                    selectors = [
                        'nav a', 'header a', 'ul.menu a', 'a', 'footer a'
                    ]
                    found_link = None
                    for txt in link_text_candidates:
                        lowered = txt.lower()
                        try:
                            found_link = WebDriverWait(driver, 5).until(
                                EC.element_to_be_clickable((By.XPATH, f"//a[contains(translate(normalize-space(.), 'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÂÊÔÃÕÇ', 'abcdefghijklmnopqrstuvwxyzáéíóúâêôãõç'), '{lowered}')]"))
                            )
                            if found_link:
                                break
                        except TimeoutException:
                            pass
                        for sel in selectors:
                            try:
                                links = driver.find_elements(By.CSS_SELECTOR, sel)
                                for lk in links:
                                    try:
                                        t = (lk.text or '').strip().lower()
                                        if lowered in t and lk.is_displayed():
                                            found_link = lk
                                            break
                                    except Exception:
                                        continue
                                if found_link:
                                    break
                            except Exception:
                                continue
                        if found_link:
                            break
                    if not found_link:
                        print(f"[auto_fill] Link não encontrado para: {link_text_candidates}")
                        return None
                    try:
                        driver.execute_script('arguments[0].click();', found_link)
                    except Exception:
                        found_link.click()
                    try:
                        WebDriverWait(driver, 15).until(
                            EC.presence_of_element_located((By.CSS_SELECTOR, 'body, html'))
                        )
                    except TimeoutException:
                        pass
                    out_path = os.path.join(home_dir, f'{out_filename_prefix}_{domain_slug}_{int(time.time())}.png')
                    driver.save_screenshot(out_path)
                    try:
                        gui_entry_widget.delete(0, tk.END)
                        gui_entry_widget.insert(0, out_path)
                    except Exception:
                        pass
                    print(f'[auto_fill] {out_filename_prefix} screenshot: {out_path}')
                    return out_path

                # Quem somos / Sobre nós
                go_and_shoot([
                    'Quem somos', 'Quem Somos', 'Sobre nós', 'Sobre', 'About'
                ], 'quem_somos', self.entry_pagina_quem_somos)

                # Contato
                go_and_shoot([
                    'Contato', 'Contact'
                ], 'contato', self.entry_pagina_contato)

                # Produtos / Loja
                go_and_shoot([
                    'Produtos', 'Loja', 'Shop', 'Catálogo'
                ], 'produtos', self.entry_pagina_produtos)

                # Volta para o painel (wp-admin)
                driver.get(dominio.rstrip('/') + '/wp-admin/')
                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'adminmenumain')))

                # Produtos -> Todos os produtos
                products_link = None
                try:
                    products_link = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, "#menu-posts-product a[href*='edit.php?post_type=product']"))
                    )
                except TimeoutException:
                    try:
                        products_link = WebDriverWait(driver, 10).until(
                            EC.element_to_be_clickable((By.XPATH, "//div[@id='adminmenumain']//a[contains(translate(normalize-space(.), 'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÂÊÔÃÕÇ', 'abcdefghijklmnopqrstuvwxyzáéíóúâêôãõç'),'produtos') or contains(@href, 'edit.php?post_type=product')]"))
                        )
                    except TimeoutException:
                        products_link = None
                if products_link:
                    driver.execute_script('arguments[0].click();', products_link)
                    try:
                        WebDriverWait(driver, 15).until(
                            EC.presence_of_element_located((By.CSS_SELECTOR, '#wpbody-content'))
                        )
                    except TimeoutException:
                        pass
                    todos_produtos_path = os.path.join(home_dir, f'todos_produtos_{domain_slug}_{int(time.time())}.png')
                    driver.save_screenshot(todos_produtos_path)
                    try:
                        self.entry_todos_produtos.delete(0, tk.END)
                        self.entry_todos_produtos.insert(0, todos_produtos_path)
                    except Exception:
                        pass
                    print(f'[auto_fill] Todos os Produtos screenshot: {todos_produtos_path}')

                # ----------------- FIM FLUXO 1 -----------------

                # Menu Astra
                astra = None
                for xp in ["//a[contains(text(), 'Astra')]", "//li[contains(@class,'menu-top')]//a[contains(text(),'Astra')]", "//div[@id='adminmenumain']//a[contains(text(),'Astra')]"]:
                    try:
                        astra = driver.find_element(By.XPATH, xp)
                        break
                    except NoSuchElementException:
                        continue
                if not astra:
                    raise RuntimeError('Menu Astra não encontrado')
                driver.execute_script('arguments[0].click();', astra)
                time.sleep(2)

                # Personalizar
                customize = None
                for xp in ["//a[contains(text(), 'Personalizar')]", "//a[contains(text(),'Customize')]", "//a[contains(@href,'customize.php')]"]:
                    try:
                        customize = driver.find_element(By.XPATH, xp)
                        break
                    except NoSuchElementException:
                        continue
                if not customize:
                    raise RuntimeError('Botão Personalizar não encontrado')
                driver.execute_script('arguments[0].click();', customize)
                WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.ID, 'customize-controls')))
                time.sleep(5)

                # Style Guide (botão com id 'astra-tour')
                print('[auto_fill] Procurando botão Style Guide...')
                style_btn = None
                try:
                    # Primeiro em default_content
                    driver.switch_to.default_content()
                    style_btn = WebDriverWait(driver, 20).until(
                        EC.element_to_be_clickable((By.ID, 'astra-tour'))
                    )
                except TimeoutException:
                    print('[auto_fill] Botão por ID não encontrado em 20s, tentando seletores alternativos...')
                    # Tentar múltiplos seletores com WebDriverWait
                    selectors = [
                        'button#astra-tour',
                        'button[name="astra-tour"]',
                        'button[title="Style Guide"]',
                        'button:has(.ast-style-guide-tooltip)',
                        '.ast-style-guide',
                        'button .ast-style-guide-tooltip',
                        'div.ast-style-guide'  # Novo: possível contêiner geral
                    ]
                    found = None
                    for sel in selectors:
                        try:
                            found = WebDriverWait(driver, 10).until(
                                EC.element_to_be_clickable((By.CSS_SELECTOR, sel))
                            )
                            if found:
                                style_btn = found
                                break
                        except TimeoutException:
                            continue
                    # Se ainda não achou, tentar dentro de iframes
                    if not style_btn:
                        print('[auto_fill] Tentando localizar dentro de iframes...')
                        try:
                            iframes = driver.find_elements(By.TAG_NAME, 'iframe')
                            for iframe in iframes:
                                try:
                                    driver.switch_to.frame(iframe)
                                    for sel in selectors:
                                        try:
                                            found = WebDriverWait(driver, 5).until(
                                                EC.element_to_be_clickable((By.CSS_SELECTOR, sel))
                                            )
                                            if found:
                                                style_btn = found
                                                break
                                        except TimeoutException:
                                            continue
                                    if style_btn:
                                        break
                                except Exception:
                                    continue
                                finally:
                                    driver.switch_to.default_content()
                        except Exception:
                            driver.switch_to.default_content()
                if not style_btn:
                    raise RuntimeError('Botão Style Guide (astra-tour) não encontrado. Verifique se o Astra está ativo.')
                print('[auto_fill] Botão encontrado, clicando...')
                driver.execute_script('arguments[0].click();', style_btn)

                # Aguardar painel Style Guide aparecer (obrigatório)
                print('[auto_fill] Aguardando painel Style Guide (div.ast-styler-card)...')
                panel_selectors = [
                    'div.ast-styler-card',
                    'div.ast-style-guide',  # Fallback para variações
                    'div[class*="style-guide"]',  # Qualquer classe com "style-guide"
                    'div[class*="styler"]'  # Qualquer classe com "styler"
                ]
                panel_found = None
                try:
                    # Primeiro no default_content
                    driver.switch_to.default_content()
                    for sel in panel_selectors:
                        try:
                            panel_found = WebDriverWait(driver, 30).until(
                                EC.presence_of_element_located((By.CSS_SELECTOR, sel))
                            )
                            print(f'[auto_fill] Painel encontrado com seletor: {sel}')
                            break
                        except TimeoutException:
                            continue
                    # Se não achou, tenta em iframes
                    if not panel_found:
                        print('[auto_fill] Painel não encontrado no default_content, tentando iframes...')
                        iframes = driver.find_elements(By.TAG_NAME, 'iframe')
                        for iframe in iframes:
                            try:
                                driver.switch_to.frame(iframe)
                                for sel in panel_selectors:
                                    try:
                                        panel_found = WebDriverWait(driver, 10).until(
                                            EC.presence_of_element_located((By.CSS_SELECTOR, sel))
                                        )
                                        print(f'[auto_fill] Painel encontrado em iframe com seletor: {sel}')
                                        break
                                    except TimeoutException:
                                        continue
                                if panel_found:
                                    break
                            except Exception:
                                continue
                            finally:
                                driver.switch_to.default_content()
                    if not panel_found:
                        # Captura o HTML atual para depuração
                        html = driver.page_source[:1000]  # Limita para não sobrecarregar
                        messagebox.showerror('Erro', f'Painel Style Guide não apareceu (ast-styler-card). HTML parcial:\n{html}')
                        raise RuntimeError('Painel Style Guide não apareceu a tempo. Verifique o tema Astra.')
                except TimeoutException:
                    html = driver.page_source[:1000]
                    messagebox.showerror('Erro', f'Painel Style Guide não apareceu (ast-styler-card). HTML parcial:\n{html}')
                    raise RuntimeError('Painel Style Guide não apareceu a tempo. Verifique o tema Astra.')

                # Screenshot do Style Guide
                print('[auto_fill] Capturando screenshot do Style Guide...')
                ok = driver.save_screenshot(out_img)
                if not ok:
                    raise RuntimeError('Falha ao salvar screenshot')
                print(f'[auto_fill] Screenshot salvo: {out_img}')

                # Fechar o customizer (voltar para o site) e tirar print da Home
                try:
                    close_btn = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, 'a.customize-controls-close'))
                    )
                except TimeoutException:
                    print('[auto_fill] Botão close não encontrado por CSS, tentando XPath...')
                    try:
                        close_btn = WebDriverWait(driver, 5).until(
                            EC.element_to_be_clickable((By.XPATH, "//a[contains(@class, 'customize-controls-close')]"))
                        )
                    except TimeoutException:
                        close_btn = None
                if close_btn:
                    print('[auto_fill] Fechando customizer...')
                    driver.execute_script('arguments[0].click();', close_btn)
                else:
                    print('[auto_fill] Botão close não encontrado, continuando...')

                # Tentar encontrar a página Home
                print('[auto_fill] Aguardando página Home...')
                home_selectors = [
                    'div.hfeed.site',
                    'div.site-container',  # Fallback para outros temas
                    'body.home',  # Fallback genérico
                    'div#page'  # Container comum no WordPress
                ]
                home_found = None
                try:
                    driver.switch_to.default_content()
                    for sel in home_selectors:
                        try:
                            home_found = WebDriverWait(driver, 10).until(
                                EC.presence_of_element_located((By.CSS_SELECTOR, sel))
                            )
                            print(f'[auto_fill] Página Home encontrada com seletor: {sel}')
                            break
                        except TimeoutException:
                            continue
                    if not home_found:
                        html = driver.page_source[:1000]
                        messagebox.showerror('Erro', f'Página Home não encontrada (hfeed site). HTML parcial:\n{html}')
                        raise RuntimeError('Página Home não encontrada. Verifique o tema ou estrutura do site.')
                except TimeoutException:
                    html = driver.page_source[:1000]
                    messagebox.showerror('Erro', f'Página Home não encontrada (hfeed site). HTML parcial:\n{html}')
                    raise RuntimeError('Página Home não encontrada. Verifique o tema ou estrutura do site.')

                # Salvar screenshot da Home
                print('[auto_fill] Capturando screenshot da Home...')
                home_dir = os.path.dirname(out_img)
                home_path = os.path.join(home_dir, os.path.basename(out_img).replace('style_guide_', 'home_'))
                driver.save_screenshot(home_path)
                try:
                    self.entry_pagina_home.delete(0, tk.END)
                    self.entry_pagina_home.insert(0, home_path)
                except Exception:
                    pass
                print(f'[auto_fill] Screenshot da Home salvo: {home_path}')

                messagebox.showinfo('Sucesso', f'Style Guide: {out_img}\nPágina Home: {home_path}')
            except (WebDriverException, TimeoutException) as e:
                messagebox.showerror('Erro', f'Falha na automação: {e}')
            except Exception as e:
                messagebox.showerror('Erro', str(e))
            finally:
                if driver:
                    driver.quit()

        def run(self):
            template = self.entry_template.get().strip()
            if not template or not Path(template).exists():
                messagebox.showerror('Erro', 'Template .docx inválido ou não informado')
                return
            cnpj_raw = self.entry_cnpj.get().strip()
            try:
                cnpj_norm = normalize_cnpj(cnpj_raw)
            except Exception as e:
                messagebox.showerror('Erro', f'CNPJ inválido: {e}')
                return
            try:
                data = consulta_empresa(cnpj_norm)
            except Exception as e:
                messagebox.showerror('Erro', f'Falha ao consultar ReceitaWS: {e}')
                return
            mapping = build_mapping(data)
            mapping["DATA_BACKUP"] = self.entry_data_backup.get().strip()
            mapping["DATA_KICKOFF"] = self.entry_data_kickoff.get().strip()
            mapping["DATA_ENTREGA"] = self.entry_data_entrega.get().strip()
            dominio = self.entry_dominio.get().strip()
            mapping["DOMINIO"] = dominio
            if dominio:
                mapping["DOMINIOWP"] = dominio + "/wp-admin/"
            else:
                mapping["DOMINIOWP"] = ""
            mapping["DEMANDA"] = self.entry_demanda.get().strip()
            mapping["EMAIL_CRIADO"] = self.entry_email_criado.get().strip()
            # Hospedagem
            hospedagem_list = []
            if self.hostinger_var.get():
                hospedagem_list.append("HOSTINGER")
            if self.localweb_var.get():
                hospedagem_list.append("LOCALWEB")
            if self.uol_var.get():
                hospedagem_list.append("UOL")
            if hospedagem_list:
                mapping["HOSPEDAGEM"] = ", ".join(hospedagem_list)
            else:
                mapping["HOSPEDAGEM"] = self.entry_hospedagem_custom.get().strip()
            mapping["IDENTIDADE_VISUAL_E_PALETA_DE_CORES"] = self.entry_identidade.get().strip()
            mapping["PAGINA_HOME_IMG"] = self.entry_pagina_home.get().strip()
            mapping["PAGINA_PRODUTOS_IMG"] = self.entry_pagina_produtos.get().strip()
            mapping["PAGINA_QUEM_SOMOS_IMG"] = self.entry_pagina_quem_somos.get().strip()
            mapping["PAGINA_CONTATO_IMG"] = self.entry_pagina_contato.get().strip()
            mapping["DETALHES_PEDIDO"] = self.entry_detalhes_pedido.get().strip()
            mapping["DETALHES_PRODUTO"] = self.entry_detalhes_produto.get().strip()
            mapping["TODOS_PRODUTOS"] = self.entry_todos_produtos.get().strip()
            mapping["ESPECIALISTARESPONSAVEL"] = "ITALO FELIPE IGNACIO"
            drive = self.entry_drive.get().strip()
            if drive:
                if not drive.startswith(('http://', 'https://')):
                    drive = 'https://' + drive
                mapping['LINK_DRIVE'] = drive
                mapping['LINK_DRIVE_TEXT'] = self.entry_drive_text.get().strip() or 'Link Drive'
                mapping['LINK_PARA_DOWNLOAD'] = drive
                mapping['LINK_PARA_DOWNLOAD_TEXT'] = 'Link para download'
            if self.use_ai_var.get():
                source_parts = []
                if mapping.get('ATIVIDADE_PRINCIPAL'):
                    source_parts.append('Atividade principal: ' + mapping['ATIVIDADE_PRINCIPAL'])
                if mapping.get('RESUMO_EMPRESA_CLIENTE'):
                    source_parts.append('Resumo: ' + mapping['RESUMO_EMPRESA_CLIENTE'])
                source_text = '\n'.join(source_parts).strip() or str(data)[:2000]
                try:
                    ai = get_ai_provider(self.ai_provider.get())
                except Exception as e:
                    messagebox.showwarning('Aviso', f'Erro ao inicializar provedor IA: {e}\nUsando mock.')
                    ai = MockProvider()
                try:
                    objetivo = ai.generate_objective(source_text, mapping)
                except Exception as e:
                    messagebox.showwarning('Aviso', f'Erro ao gerar objetivo com IA: {e}\nUsando heurística local.')
                    objetivo = MockProvider().generate_objective(source_text, mapping)
                mapping['OBJETIVO_EMPRESA'] = objetivo
            else:
                mapping['OBJETIVO_EMPRESA'] = ''
            out = self.entry_out.get().strip() or f'relatorio_{cnpj_norm}.docx'
            try:
                process_document(template, out, mapping)
            except Exception as e:
                messagebox.showerror('Erro', f'Erro ao processar documento: {e}')
                return
            messagebox.showinfo('Sucesso', f'Relatório gerado: {out}')

# ----------------- Main -----------------
def main():
    fix_docx_templates()
    parser = argparse.ArgumentParser(description="Preencher relatórios Word via CNPJ")
    parser.add_argument("--template", help=".docx template")
    parser.add_argument("--cnpj", help="CNPJ da empresa")
    parser.add_argument("--drive", help="Link do Drive")
    parser.add_argument("--drive-text", help="Texto do link do Drive")
    parser.add_argument("--use-ai", action="store_true", help="Usar IA para preencher [OBJETIVO_EMPRESA]")
    parser.add_argument("--ai-provider", help="Provedor IA: mock/hf/openai")
    parser.add_argument("--out", help="Arquivo de saída .docx")
    parser.add_argument("--data-backup", help="Data de backup [DATA_BACKUP]")
    parser.add_argument("--data-kickoff", help="Data de kickoff [DATA_KICKOFF]")
    parser.add_argument("--data-entrega", help="Data de entrega [DATA_ENTREGA]")
    parser.add_argument("--dominio", help="Texto [DOMINIO]")
    parser.add_argument("--demanda", help="Texto [DEMANDA]")
    parser.add_argument("--hospedagem", help="Texto [HOSPEDAGEM]")
    parser.add_argument("--email-criado", help="Texto [EMAIL_CRIADO]")
    parser.add_argument("--identidade-visual", help="Caminho da imagem [IDENTIDADE_VISUAL_E_PALETA_DE_CORES]")
    parser.add_argument("--pagina-home-img", help="Caminho da imagem [PAGINA_HOME_IMG]")
    parser.add_argument("--pagina-produtos-img", help="Caminho da imagem [PAGINA_PRODUTOS_IMG]")
    parser.add_argument("--pagina-quem-somos-img", help="Caminho da imagem [PAGINA_QUEM_SOMOS_IMG]")
    parser.add_argument("--pagina-contato-img", help="Caminho da imagem [PAGINA_CONTATO_IMG]")
    parser.add_argument("--detalhes-pedido", help="Caminho da imagem [DETALHES_PEDIDO]")
    parser.add_argument("--detalhes-produto", help="Caminho da imagem [DETALHES_PRODUTO]")
    parser.add_argument("--todos-produtos", help="Caminho da imagem [TODOS_PRODUTOS]")
    parser.add_argument("--run-tests", action="store_true", help="Executar testes rápidos")
    args = parser.parse_args()
    if args.run_tests:
        print("Testes simples:")
        try:
            assert normalize_cnpj("12.345.678/0001-95") == "12345678000195"
            assert normalize_cnpj("12345678000195") == "12345678000195"
            print("normalize_cnpj OK")
        except AssertionError:
            print("normalize_cnpj falhou")
        try:
            doc = Document()
            p = doc.add_paragraph("[NOME_EMPRESA_CLIENTE] e [CNPJ]")
            mapping = {"NOME_EMPRESA_CLIENTE":"ACME","CNPJ":"123"}
            replace_in_paragraph(p, mapping)
            assert "ACME" in p.text and "123" in p.text
            print("replace_in_paragraph OK")
        except Exception as e:
            print("replace_in_paragraph falhou:", e)
        # Teste da normalização de imagem
        try:
            test_path = "/path with spaces/image.jpg"
            normalized = normalize_image_path(test_path)
            print(f"normalize_image_path: '{test_path}' -> '{normalized}' OK")
        except Exception as e:
            print("normalize_image_path falhou:", e)
        return
    if TKINTER_AVAILABLE and not any([args.template, args.cnpj, args.drive]):
        root = tk.Tk()
        app = App(root)
        root.mainloop()
    else:
        extra_mapping = {
            "DATA_BACKUP": args.data_backup or "",
            "DATA_KICKOFF": args.data_kickoff or "",
            "DATA_ENTREGA": args.data_entrega or "",
            "DOMINIO": args.dominio or "",
            "DEMANDA": args.demanda or "",
            "HOSPEDAGEM": args.hospedagem or "",
            "EMAIL_CRIADO": args.email_criado or "",
            "IDENTIDADE_VISUAL_E_PALETA_DE_CORES": args.identidade_visual or "",
            "PAGINA_HOME_IMG": args.pagina_home_img or "",
            "PAGINA_PRODUTOS_IMG": args.pagina_produtos_img or "",
            "PAGINA_QUEM_SOMOS_IMG": args.pagina_quem_somos_img or "",
            "PAGINA_CONTATO_IMG": args.pagina_contato_img or "",
            "DETALHES_PEDIDO": args.detalhes_pedido or "",
            "DETALHES_PRODUTO": args.detalhes_produto or "",
            "TODOS_PRODUTOS": args.todos_produtos or "",
        }
        run_cli(
            template=args.template,
            cnpj=args.cnpj,
            drive=args.drive,
            drive_text=args.drive_text,
            use_ai=args.use_ai,
            ai_provider=args.ai_provider,
            out=args.out,
            extra_mapping=extra_mapping
        )

if __name__ == "__main__":
    main()
